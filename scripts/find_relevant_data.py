"""
Script to search for docx, pdf, and txt files in a directory and its
subdirectories.  It uses chatGPT to look at the filename and first page of the
document to determine if it is relevant, then saves the filepath and relevance
to a json file in the root directory of the SUNY school.
"""

import os
import sys
import docx
import json
import PyPDF2

from tqdm import tqdm
from openai import OpenAI

from src.agent import Agent
from src.utils import parse_json
from src.pdf_tools import load_pdf_text
from src.prompts import FILTER_FILES_PROMPT

opj = os.path.join

def find_files(data_dir):
    """
    Find the full paths to all PDFs, documents, and images larger than 200x200 in the specified directory.

    Args:
        data_dir (str): The directory to search for files.

    Returns:
        list[str]: A list of full paths to PDFs, documents, and images larger than 200x200.
    """
    file_paths = []
    for root, dirs, files in os.walk(data_dir):
        for file in files:
            if file.lower().endswith(('.pdf', '.doc', '.docx', '.txt')):
                file_path = os.path.join(root, file)
                file_paths.append(file_path)
    return file_paths


def get_pdf_metadata(file_path):
    with open(file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        return str(pdf_reader.metadata)


def get_docx_metadata(file_path):
    try:
        doc = docx.Document(file_path)
        core_properties = doc.core_properties
        metadata = {
            'author': core_properties.author,
            'title': core_properties.title,
            'subject': core_properties.subject,
            'created': str(core_properties.created) if core_properties.created else None,
            'modified': str(core_properties.modified) if core_properties.modified else None,
            'last_modified_by': core_properties.last_modified_by,
            'keywords': core_properties.keywords,
            'category': core_properties.category,
            'comments': core_properties.comments,
        }
        return str(metadata)
    except Exception as e:
        return f"Error reading DOCX metadata: {str(e)}"


def print_response(response):
    filepath = response['filepath']
    relevant = response['relevant'].upper()
    max_length = max(len('Filepath: ' +filepath), len('Relevant: ' +relevant), 10)  # Minimum width of 10

    print('\n')
    print('----------------------------------------')
    print('Filepath: ' + filepath)
    print('Relevant: ' + relevant)
    print('----------------------------------------')


def read_docx(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def main():

    dirs_to_omit = [
        'system_data/suny_schools/State-University-Of-New-York-College-Of-Technology-At-Canton/www.canton.edu/media/curriculum/'
    ]

    data_dir = sys.argv[1]

    relevant_files_path = opj(data_dir, 'relevant_files.json')
    if os.path.exists(relevant_files_path):
        with open(relevant_files_path, 'r') as f:
            relevant_files = json.load(f)
    else:
        relevant_files = {}

    print('Existing relevant files:', len(relevant_files))

    file_paths = find_files(data_dir)
    print('Total files:', len(file_paths))

    file_paths = [fp for fp in file_paths if fp not in relevant_files]
    print('Total files to process before omitting:', len(file_paths))

    file_paths = [fp for fp in file_paths if dirs_to_omit[0] not in fp]
    print('Total files to process after omitting:', len(file_paths))

    import random
    random.seed(523)
    random.shuffle(file_paths)

    client = OpenAI(api_key=os.getenv("PATHFINDER_OPENAI_API_KEY"))
    system_prompt = 'Your task is to determine if the given document is relevant to a high school student who is thinking about applying to the university.'
    agent = Agent(client, 'Relevance Detector', None, system_prompt, model='gpt-4o-mini', json_mode=True)

    for idx, filepath in tqdm(enumerate(file_paths), total=len(file_paths), desc="Processing files"):

        extension = os.path.splitext(filepath)[1]

        try:
            if filepath in relevant_files:
                print(f'Skipping {filepath} because it has already been processed.')
                continue

            prompt = FILTER_FILES_PROMPT + '\n'
            prompt += f'Document Filename: {filepath}\n'
            if extension == '.pdf':
                prompt += f'Document Metadata: {get_pdf_metadata(filepath)}\n'
            elif extension == '.docx':
                prompt += f'Document Metadata: {get_docx_metadata(filepath)}\n'
            prompt += 'Document Content:\n'

            if extension == '.pdf':
                document_content = '\n'.join(load_pdf_text(filepath))
            elif extension == '.docx':
                document_content = read_docx(filepath)
            elif extension == '.txt':
                document_content = open(filepath, 'r').read()

            prompt += document_content

            agent.add_message('user', prompt)
            response = parse_json(agent.invoke().choices[0].message.content)
            agent.delete_last_message()
            filepath = response['filepath']
            relevant = response['relevant'].upper()
            #if relevant == 'YES':
            #    print(prompt)
            #    print('-'*100)
            #    print(response)
            #    input('Press Enter to continue...')

            relevant_files[filepath] = relevant

        except Exception as e:
            print(f'Error processing {filepath}: {e}')
            continue

        if not idx % 10 and idx > 0:
            with open(relevant_files_path, 'w') as f:
                json.dump(relevant_files, f)
            print('Saved to', relevant_files_path)

    with open(relevant_files_path, 'w') as f:
        json.dump(relevant_files, f)
    print('Saved to', relevant_files_path)

if __name__ == "__main__":
    main()